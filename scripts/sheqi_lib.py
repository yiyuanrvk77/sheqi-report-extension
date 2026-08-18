# -*- coding: utf-8 -*-
"""
sheqi-report 核心库：主体管理 / 记录导入 / 验证码人工输入 / 12377 自动填表提交。
所有数据默认保存在 <cwd>/sheqi_data，可用环境变量 SHEQI_DATA_DIR 覆盖。
"""
import csv
import io
import json
import os
import re
import shutil
import subprocess
import sys
import time
from datetime import datetime

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

DEFAULT_REPORT_URL = "https://www.12377.cn/jbxzxq/sql_web.html"
# 注意：不再设本地固定每日上限（实际以服务端“今日举报次数已达上限”为准），
# 如需本地软限制可在 run 时用 --daily N 指定。
SUBMIT_INTERVAL = 15
MAX_CAPTCHA_RETRY = 4
CONTENT_MAX = 500

MAGIC_EXT = [
    (b"\xff\xd8\xff", "jpg"),
    (b"\x89PNG\r\n\x1a\n", "png"),
    (b"%PDF-", "pdf"),
    (b"GIF8", "gif"),
]


def data_dir():
    d = os.environ.get("SHEQI_DATA_DIR")
    return os.path.abspath(d) if d else os.path.join(os.getcwd(), "sheqi_data")


def ensure_dirs():
    for sub in ("", "documents", "output", "logs"):
        os.makedirs(os.path.join(data_dir(), sub), exist_ok=True)


def now_str():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def today():
    return datetime.now().strftime("%Y-%m-%d")


def truncate(s, n):
    s = str(s or "").strip()
    if len(s) <= n:
        return s
    if n <= 3:
        return s[:n]
    return s[: n - 3].rstrip() + "..."


def load_json(path, default):
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8-sig") as f:
                return json.load(f)
    except Exception:
        pass
    return default


def save_json(path, obj):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


SUBJECTS_FILE = os.path.join(data_dir(), "subjects.json")
RECORDS_FILE = os.path.join(data_dir(), "records.json")
RESULTS_FILE = os.path.join(data_dir(), "results.json")


# ================= 多主体管理 =================
def load_subjects():
    return load_json(SUBJECTS_FILE, {"current": None, "subjects": {}})


def save_subjects(data):
    save_json(SUBJECTS_FILE, data)


def subject_names():
    return list(load_subjects().get("subjects", {}).keys())


def get_subject(name=None):
    data = load_subjects()
    subs = data.get("subjects", {})
    if not subs:
        return None
    if name:
        return subs.get(name)
    return subs.get(data.get("current")) or next(iter(subs.values()))


def set_current(name):
    data = load_subjects()
    if name not in data.get("subjects", {}):
        raise ValueError("主体不存在: %s（现有: %s）" % (name, ", ".join(subject_names()) or "无"))
    data["current"] = name
    save_subjects(data)
    return name


def detect_ext(path):
    try:
        with open(path, "rb") as f:
            head = f.read(8)
        for magic, ext in MAGIC_EXT:
            if head.startswith(magic):
                return ext
    except Exception:
        pass
    return os.path.splitext(path)[1].lstrip(".") or "bin"


def register_file(subject_name, key, src):
    """把材料复制到 documents/<主体>/<键>.<真实扩展名>，自动修正扩展名。"""
    if not src or not os.path.exists(src):
        return None
    ext = detect_ext(src)
    target_dir = os.path.join(data_dir(), "documents", subject_name)
    os.makedirs(target_dir, exist_ok=True)
    target = os.path.join(target_dir, "%s.%s" % (key, ext))
    shutil.copy2(src, target)
    return os.path.relpath(target, data_dir()).replace("\\", "/")


def add_subject(name, fields=None, files=None):
    data = load_subjects()
    subs = data.setdefault("subjects", {})
    sub = subs.get(name, {})
    sub["name"] = name
    for k, v in (fields or {}).items():
        if v is not None:
            sub[k] = v
    fmap = sub.setdefault("files", {})
    for k, src in (files or {}).items():
        if src:
            rel = register_file(name, k, src)
            if rel:
                fmap[k] = rel
    subs[name] = sub
    if not data.get("current"):
        data["current"] = name
    save_subjects(data)
    return sub


def remove_subject(name):
    data = load_subjects()
    subs = data.get("subjects", {})
    if name not in subs:
        return False
    del subs[name]
    if data.get("current") == name:
        data["current"] = next(iter(subs)) if subs else None
    save_subjects(data)
    return True


def resolve_doc(subject, key):
    """返回材料绝对路径；优先按相对 data_dir 路径解析，再回退绝对路径与 documents/。"""
    files = subject.get("files") or {}
    rel = files.get(key)
    if not rel:
        return None
    # register_file 存的是 relpath(target, data_dir)，优先拼 data_dir 命中（避免误解析 cwd 同名文件）
    p = os.path.join(data_dir(), rel)
    if not os.path.exists(p):
        p = os.path.abspath(rel)
    if not os.path.exists(p):
        p = os.path.join(data_dir(), "documents", subject.get("name", ""), os.path.basename(rel))
    if not os.path.exists(p):
        p = os.path.join(data_dir(), "documents", os.path.basename(rel))
    return p if os.path.exists(p) else None


def subject_missing_files(subject):
    return [k for k in ("license", "id_front", "id_back", "auth") if not resolve_doc(subject, k)]


# ================= 记录存储 =================
def load_records():
    return load_json(RECORDS_FILE, [])


def save_records(records):
    save_json(RECORDS_FILE, records)


def load_results():
    return load_json(RESULTS_FILE, [])


def save_results(results):
    save_json(RESULTS_FILE, results)


def pending_records(records, subject=None):
    out = []
    for r in records:
        if r.get("status") in ("pending", "failed") and (not subject or r.get("subject") == subject):
            out.append(r)
    return out


def recover_stuck(subject=None):
    """将状态卡在 submitting 的记录回滚为 pending（run 中断/崩溃后使用）。"""
    records = load_records()
    n = 0
    for r in records:
        if r.get("status") == "submitting" and (not subject or r.get("subject") == subject):
            r["status"] = "pending"
            r["error"] = "recovered"
            n += 1
    if n:
        save_records(records)
        print("已恢复 %d 条卡死的 submitting 记录为 pending。" % n)
    else:
        print("没有卡死的 submitting 记录。")
    return n


def count_today_success(results, subject=None):
    d = today()
    n = 0
    for r in results:
        if r.get("success") and str(r.get("date", ""))[:10] == d:
            if not subject or r.get("subject") == subject:
                n += 1
    return n


def export_csv(records):
    fields = ["id", "subject", "date", "platform", "account", "url", "title",
              "content", "status", "query_code", "error", "submit_time"]
    path = os.path.join(data_dir(), "records.csv")
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        for r in records:
            w.writerow(r)
    return path


def export_results_excel(results, note=None):
    """导出提交结果 Excel（含查询码），保存到 sheqi_data/results.xlsx。
    note 非空时在表格末尾追加一行说明（如“今日举报次数已达上限”）。"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
    except Exception:
        return None
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "提交结果"
    headers = ["提交时间", "ID", "主体", "平台", "网址", "状态", "查询码", "错误信息"]
    ws.append(headers)
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="2F5496")
        c.alignment = Alignment(horizontal="center")
    for r in results:
        ws.append([
            r.get("time", ""), r.get("id", ""), r.get("subject", ""),
            r.get("platform", ""), r.get("url", ""),
            "成功" if r.get("success") else "失败",
            r.get("code", ""), r.get("error", "") or "",
        ])
    if note:
        ws.append([note])
        last = ws.max_row
        ws.merge_cells(start_row=last, start_column=1, end_row=last, end_column=8)
        c = ws.cell(row=last, column=1)
        c.font = Font(bold=True, color="C00000")
        c.alignment = Alignment(horizontal="left")
    widths = [20, 22, 10, 14, 50, 8, 24, 40]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    path = os.path.join(data_dir(), "results.xlsx")
    wb.save(path)
    return path


def export_today_excel(subject=None, date=None):
    """导出指定日期（默认今天）已提交且含查询码的记录为 Excel（最后一步归档）。
    筛选 records.json 中 status=submitted 且 query_code 非空的记录，
    保存到 sheqi_data/reports/今日提交_YYYY-MM-DD.xlsx，返回 (path, 条数)。
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    date = date or today()
    records = load_records()
    rows = [
        r for r in records
        if r.get("status") == "submitted"
        and r.get("query_code")
        and (not subject or r.get("subject") == subject)
        and str(r.get("submit_time") or "").startswith(date)
    ]
    # 兼容 submit_time 缺失/不同格式：回退按 query_code 且导入日期判断
    if not rows:
        rows = [
            r for r in records
            if r.get("status") == "submitted"
            and r.get("query_code")
            and (not subject or r.get("subject") == subject)
        ]
    # 按提交时间排序（无则按 ID）
    rows.sort(key=lambda r: str(r.get("submit_time") or r.get("id") or ""))
    if not rows:
        return None, 0

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "今日提交"
    headers = ["序号", "记录ID", "审核码/查询码", "平台", "账号名称", "发布日期", "标题", "网址", "提交时间"]
    ws.append(headers)
    head_fill = PatternFill("solid", fgColor="2F5496")
    thin = Border(*[Side(style="thin", color="D9D9D9")] * 4)
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF", size=11)
        c.fill = head_fill
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = thin
    for i, r in enumerate(rows, 1):
        url = r.get("url") or ""
        ws.append([
            i, r.get("id", ""), r.get("query_code", ""), r.get("platform", ""),
            r.get("account", ""), r.get("date", ""), r.get("title", ""), url,
            r.get("submit_time", ""),
        ])
        cell = ws.cell(row=i + 1, column=8)
        if str(url).startswith("http"):
            cell.hyperlink = url
            cell.font = Font(color="0563C1", underline="single")
        for c in ws[i + 1]:
            c.border = thin
        qc = ws.cell(row=i + 1, column=3)
        qc.font = Font(bold=True, color="C00000")
    widths = [6, 18, 26, 14, 18, 12, 46, 46, 20]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    # 汇总行
    total = ws.max_row + 1
    ws.cell(row=total, column=1, value="合计").font = Font(bold=True)
    ws.merge_cells(start_row=total, start_column=1, end_row=total, end_column=2)
    c = ws.cell(row=total, column=3, value="%d 条" % len(rows))
    c.font = Font(bold=True, color="C00000")

    out_dir = os.path.join(data_dir(), "reports")
    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, "今日提交_%s.xlsx" % date)
    wb.save(path)
    return path, len(rows)
# ================= 导入：自动识别格式 =================
FIELD_ALIASES = {
    "url": ["链接", "网址", "url", "link", "举报网址", "原始链接", "原文链接", "原文地址", "稿件链接",
            "新闻链接", "帖子地址", "网页链接", "地址", "文章链接", "网络链接", "帖子链接", "源链接"],
    "title": ["标题", "title", "文章标题", "标题名称", "新闻标题", "文章名", "题目"],
    "body": ["正文", "内容", "摘要", "描述", "body", "content", "正文内容", "文章内容",
             "内容摘要", "文章摘要", "正文摘要"],
    "media": ["媒体名称", "平台名称", "媒体", "来源", "平台", "media", "source", "网站名称",
              "自媒体名称", "发布媒体", "新闻来源", "媒体来源", "发布平台", "所属平台", "平台类型"],
    "author": ["作者", "作者名称", "发布者", "博主", "author", "发布人", "作者名"],
    "account": ["账号", "账号名称", "平台账号", "account", "媒体账号"],
    "date": ["日期", "发布时间", "发布日期", "时间", "date", "发表日期", "接收时间", "接收时间（火山）"],
    "status": ["状态", "链接状态", "内容状态", "status", "是否转发"],
}
POSITIONAL = ["date", "title", "body", "media", "url", "author", "status"]
URL_RE = re.compile(r"^https?://", re.I)
SKIP_STATUS = ("转发", "已删除", "已失效")


def norm_key(s):
    return re.sub(r"\s+", "", str(s or "")).lower()


def header_row_score(row):
    norm = [norm_key(c) for c in row]
    score = 0
    for field, aliases in FIELD_ALIASES.items():
        for n in norm:
            if any(norm_key(a) == n for a in aliases):
                score += 1
                break
    return score


def locate_header(rows):
    """在表格前若干行里找真正的表头（兼容标题行/说明行在表头上方）。"""
    best, best_score = 0, -1
    for i, row in enumerate(rows[:12]):
        s = header_row_score([str(c or "").strip() for c in row])
        if s > best_score:
            best, best_score = i, s
    return best if best_score >= 2 else 0


def find_col(header):
    cols = {}
    norm = {i: norm_key(h) for i, h in enumerate(header)}
    for field, aliases in FIELD_ALIASES.items():
        for i, n in norm.items():
            if any(norm_key(a) == n for a in aliases):
                cols[field] = i
                break
    if "url" not in cols and len(header) == len(POSITIONAL):
        cols = {}
        for i, field in enumerate(POSITIONAL):
            cols[field] = i
    return cols


def read_text_bytes(path):
    raw = open(path, "rb").read()
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace")


def sniff_text(text):
    t = text.strip()
    if not t:
        raise ValueError("输入为空")
    if t[:1] in ("{", "["):
        return "json", None, None
    sample = t[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
        rows = list(csv.reader(io.StringIO(t), dialect))
    except Exception:
        rows = [r.split(",") for r in t.splitlines()]
    rows = [r for r in rows if any(str(c or "").strip() for c in r)]
    if not rows:
        raise ValueError("未解析到表格内容")
    h = locate_header(rows)
    header = [str(c or "").strip() for c in rows[h]]
    return "csv", header, rows[h + 1:]


def sniff_table(src, is_path):
    """返回 (fmt, header, rows)。fmt: json/csv/xlsx；自动定位表头行、多 sheet 择优。"""
    if is_path:
        if not os.path.exists(src):
            raise FileNotFoundError(src)
        with open(src, "rb") as f:
            head = f.read(8)
        if head.startswith(b"PK"):
            import openpyxl
            wb = openpyxl.load_workbook(src, data_only=True, read_only=True)
            best_rows, best_score = None, -1
            for ws in wb.worksheets:
                rows = [list(r) for r in ws.iter_rows(values_only=True)]
                rows = [[str(c or "").strip() for c in r] for r in rows if any(r)]
                if len(rows) < 2:
                    continue
                h = locate_header(rows)
                score = header_row_score(rows[h])
                if score > best_score:
                    best_score, best_rows = score, rows
            wb.close()
            if best_rows is None:
                raise ValueError("Excel 中未找到可识别的表头（需要含“链接/网址”等列名）")
            h = locate_header(best_rows)
            return "xlsx", best_rows[h], best_rows[h + 1:]
        return sniff_text(read_text_bytes(src))
    return sniff_text(src)


def normalize_json_records(obj):
    if isinstance(obj, dict):
        recs = obj.get("records") or obj.get("data") or []
    elif isinstance(obj, list):
        recs = obj
    else:
        recs = []
    out = []
    for r in recs:
        if not isinstance(r, dict):
            continue
        media = r.get("media") or r.get("platform") or ""
        out.append({
            "url": r.get("url") or r.get("link") or "",
            "title": r.get("title") or "",
            "body": r.get("content") or r.get("body") or "",
            "media": media,
            "platform": r.get("platform") or media,
            "platform_type": r.get("platform_type") or "",
            "author": r.get("author") or "",
            "account": r.get("account") or r.get("author") or media,
            "date": r.get("date") or "",
            "status": r.get("source_status") or r.get("status") or "",
            "evidence_type": r.get("evidence_type") or "P008",
            "infringement_type": r.get("infringement_type") or "1",
            "content": r.get("content") or "",
        })
    return out


def strip_html(s):
    """剥离 HTML 标签并压缩空白，保留文本内容（用于举报内容生成）"""
    if not s:
        return ""
    import re as _re
    t = _re.sub(r"<[^>]+>", "", str(s))
    t = _re.sub(r"&nbsp;|&#160;", " ", t)
    t = _re.sub(r"&amp;", "&", t)
    t = _re.sub(r"&lt;", "<", t)
    t = _re.sub(r"&gt;", ">", t)
    t = _re.sub(r"&quot;", '"', t)
    t = _re.sub(r"\s+", " ", t)
    return t.strip()


def tail_sentence(subject_name, nature=None, content_tail=None, lead="该内容"):
    """生成举报内容统一侵权定性尾句。
    content_tail 优先；否则按企业性质选内置模板（上市/拟上市用资本市场表述，其余用商业信誉表述）。
    lead: 单条内容用“该内容”，合并内容用“上述内容”。"""
    if content_tail:
        return str(content_tail).strip()
    if nature in ("SHANGSHIQIYE", "NISHANGSHIQIYE"):
        return ("%s将%s品牌声誉与公司经营、股价等无因果关系事项强行捆绑解读，"
                "片面渲染、夸大企业负面，严重损害商业信誉与资本市场舆论环境，"
                "违反《网络信息内容生态治理规定》等规定，恳请平台依法处置。"
                % (lead, subject_name))
    return ("%s严重失实，恶意贬损%s商业信誉与商品声誉，损害企业合法权益，"
            "违反《网络信息内容生态治理规定》等规定，恳请平台依法处置。"
            % (lead, subject_name))


def build_content(subject_name, date, platform, account, title, body, nature=None, content_tail=None):
    """自动整合举报内容（≤500 字）：仅统一侵权定性通用文案，
    不罗列日期/平台/账号/标题（避免机械堆叠）。
    nature: 企业性质（SHANGSHIQIYE/NISHANGSHIQIYE 等），决定定性模板；
    content_tail: 主体自定义定性文案（优先于内置模板）。"""
    return tail_sentence(subject_name, nature, content_tail, lead="该内容")


def group_records_by_date(records, url_max=500):
    """按日期降序（新→旧）分组；同一天内按 URL 长度升序（短→长）排列，
    贪心合并：从最短开始尽量往链接框塞，直到 URL 拼接总长超过 url_max(默认500) 字符即封组。
    不跨日期：同一天塞满后另起下一组，绝不打散到下一天。
    返回 [(date, [rec, ...]), ...]"""
    # 按日期分组
    by_date = {}
    for r in records:
        d = str(r.get("date") or "")
        by_date.setdefault(d, []).append(r)
    groups = []
    for date in sorted(by_date.keys(), reverse=True):
        # 同一天内按 URL 长度升序（短→长），长度相同按 id 稳定排序
        day_recs = sorted(
            by_date[date],
            key=lambda x: (len(str(x.get("url") or "")), str(x.get("id") or "")),
        )
        chunk, total = [], 0
        for r in day_recs:
            # +1 换行符（沿用原口径，略保守留 1 字符余量）
            l = len(str(r.get("url") or "")) + 1
            if chunk and total + l > url_max:
                groups.append((date, chunk))
                chunk, total = [], 0
            chunk.append(r)
            total += l
        if chunk:
            groups.append((date, chunk))
    return groups


def build_merged_content(subject_name, date, recs, nature=None, content_tail=None):
    """合并多条记录生成一条举报内容（≤500 字）。
    仅统一侵权定性通用文案（不逐条罗列日期/平台/账号/标题），URL 已在 targetUrl 中列明。"""
    return tail_sentence(subject_name, nature, content_tail, lead="上述内容")


def fit_names_to_limit(items, limit=50, sep="、"):
    """名称填满策略：去重拼接后若超 limit，从尾部逐个删除名称直到 ≤limit；
    只剩一个仍超长则截断到 limit。返回最终名称字符串。"""
    seen, ordered = [], []
    for it in items:
        s = str(it or "").strip("、，, \t")
        if s and s not in seen:
            seen.append(s)
            ordered.append(s)
    if not ordered:
        return ""
    full = sep.join(ordered)
    while len(full) > limit and len(ordered) > 1:
        ordered.pop()  # 删除最后一个名称（连顿号一起缩短）
        full = sep.join(ordered)
    if len(full) > limit:  # 只剩一个仍超长 → 截断到 limit
        return truncate(ordered[0], limit)
    return full


def resolve_content(rec, content_file=None):
    """确定具体举报内容（≤500 字）。
    优先级：1) 用户模板文件（content_file 或 sheqi_data/content_template.txt 非空）
             2) rec 中已有的 content（自动整合结果 / 导入时正文）
    返回 (content, source)：source ∈ {template, auto}"""
    # 1) 用户模板
    path = None
    if content_file:
        path = content_file
    else:
        try:
            path = os.path.join(data_dir(), "content_template.txt")
        except Exception:
            path = None
    if path and os.path.exists(path):
        try:
            with open(path, encoding="utf-8-sig") as f:
                tpl = f.read().strip()
            if tpl:
                return truncate(tpl, CONTENT_MAX), "template"
        except Exception:
            pass
    # 2) 自动整合（build_merged_content / build_content 已生成）
    return truncate(str(rec.get("content") or ""), CONTENT_MAX), "auto"


def merge_records_to_queue(records, subject, url_max=500, content_file=None):
    """把 pending 记录按最新日期排序（新→旧）、同一天内按 URL 长度升序（短→长），
    贪心合并为投递队列：从最短开始尽量往链接框塞，URL 拼接总长超过 url_max(500) 即封组，不跨日期。
    每组生成一个合并 rec：urls=多URL列表、content=合并内容、id=首条id+“+N”，
    附 platforms/accounts 去重列表（供 fill_form 名称填满策略使用）。
    返回 (queue, group_map)，group_map: 合并id -> [原始记录id, ...]"""
    sname = subject.get("name") or ""
    nature = subject.get("company_nature")
    content_tail = subject.get("content_tail")
    groups = group_records_by_date(records, url_max=url_max)
    queue, group_map = [], {}
    for date, recs in groups:
        first = recs[0]
        merged_id = "%s+%d" % (first.get("id"), len(recs))
        mrec = dict(first)
        mrec["id"] = merged_id
        mrec["date"] = date
        mrec["urls"] = [str(r.get("url") or "").strip() for r in recs if str(r.get("url") or "").startswith(("http://", "https://"))]
        mrec["url"] = mrec["urls"][0] if mrec["urls"] else first.get("url")
        # 名称填满策略：组内平台/账号去重列表（清理首尾顿号等标点，防止源数据带顿号导致双顿号）
        _TRIM = "、，, \t"
        mrec["platforms"] = [str(r.get("platform") or "").strip(_TRIM) for r in recs]
        mrec["accounts"] = [str(r.get("account") or r.get("platform") or "").strip(_TRIM) for r in recs]
        mrec["content"] = build_merged_content(sname, date, recs, nature=nature, content_tail=content_tail)
        # 用户模板优先：模板文件存在且非空则覆盖 content
        if content_file or os.path.exists(os.path.join(data_dir(), "content_template.txt")):
            mrec["content"], _src = resolve_content(mrec, content_file=content_file)
        mrec["merged_ids"] = [r.get("id") for r in recs]
        queue.append(mrec)
        group_map[merged_id] = [r.get("id") for r in recs]
    return queue, group_map


def import_records(src, subject_name, limit=0, is_path=True):
    ensure_dirs()
    subject = get_subject(subject_name)
    if not subject:
        raise ValueError("请先添加并选择主体: python sheqi.py subject add --name <名称>")
    sname = subject.get("name") or subject_name
    fmt, header, rows = sniff_table(src, is_path)
    cols = None

    if fmt == "json":
        obj = json.loads(src if not is_path else read_text_bytes(src))
        recs = normalize_json_records(obj)
    else:
        cols = find_col(header)
        if "url" not in cols:
            raise ValueError("未识别到“链接/网址”列（表格可用列: %s）" % " / ".join(str(h) for h in (header or [])))
        recs = []
        for r in rows:
            if not any(r):
                continue
            d = {}
            for field, i in cols.items():
                d[field] = r[i] if i < len(r) and r[i] is not None else ""
            recs.append(d)

    existing = load_records()
    seen = {norm_key(x.get("url")) for x in existing}
    included, skipped = [], []
    for d in recs:
        url = str(d.get("url") or "").strip()
        if not URL_RE.match(url):
            skipped.append((url, "链接非 http(s)"))
            continue
        status = str(d.get("status") or "").strip()
        if status and any(k in status for k in SKIP_STATUS):
            skipped.append((url, "状态[%s]不适合举报" % status))
            continue
        key = norm_key(url.rstrip("/"))
        if key in seen:
            skipped.append((url, "重复链接"))
            continue
        seen.add(key)
        platform = str(d.get("platform") or d.get("media") or "未知平台").strip()
        account = str(d.get("account") or d.get("author") or platform).strip()
        title = str(d.get("title") or "").strip()
        body = str(d.get("body") or "").strip()
        content = str(d.get("content") or "").strip()
        if not content:
            content = build_content(sname, d.get("date"), platform, account, title, body,
                                    nature=subject.get("company_nature"),
                                    content_tail=subject.get("content_tail"))
        included.append({
            "id": None,
            "subject": sname,
            "date": str(d.get("date") or today())[:10],
            "platform": truncate(platform, 50),
            "platform_type": str(d.get("platform_type") or "").strip(),
            "account": truncate(account, 50),
            "url": url,
            "title": truncate(title or url, 200),
            "body": body,
            "content": truncate(content, CONTENT_MAX),
            "infringement_type": str(d.get("infringement_type") or "1"),
            "evidence_type": str(d.get("evidence_type") or "P008"),
            "status": "pending",
            "query_code": "",
            "error": "",
            "submit_time": "",
            "retry_count": 0,
            "source_status": status or "未标注",
            "imported_at": now_str(),
        })

    if limit:
        included = included[:limit]
    seq_base = len(existing) + 1
    stamp = datetime.now().strftime("%Y%m%d")
    for i, rec in enumerate(included):
        rec["id"] = "CMP-%s-%04d" % (stamp, seq_base + i)
    existing.extend(included)
    save_records(existing)
    export_csv(existing)
    return {"fmt": fmt, "included": len(included), "skipped": skipped, "header": header, "cols": cols if fmt != "json" else None}
# ================= Playwright 填表提交 =================
# 12377.cn 有 WAF JS 反爬挑战（瑞数/加速乐类），需伪装浏览器特征才能进入真实页面
STEALTH_JS = """
Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
Object.defineProperty(navigator, 'languages', {get: () => ['zh-CN', 'zh', 'en']});
Object.defineProperty(navigator, 'plugins', {get: () => [1, 2, 3, 4, 5]});
window.chrome = {runtime: {}};
Object.defineProperty(navigator, 'maxTouchPoints', {get: () => 1});
const _q = window.navigator.permissions.query;
window.navigator.permissions.query = (p) => (
  p.name === 'notifications' ? Promise.resolve({state: Notification.permission}) : _q(p)
);
"""
STEALTH_UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
              "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36 Edg/126.0.0.0")


def launch_browser(headless=True):
    from playwright.sync_api import sync_playwright
    p = sync_playwright().start()
    last_err = None
    for channel in ("msedge", "chrome", None):
        try:
            if channel:
                browser = p.chromium.launch(channel=channel, headless=headless)
            else:
                browser = p.chromium.launch(headless=headless)
            return p, browser
        except Exception as e:
            last_err = e
    p.stop()
    raise RuntimeError("无法启动浏览器（Edge/Chrome 均不可用）: %s" % last_err)


# ================= 验证码 OCR（PaddleOCR 在线 API） =================
# 识别：PaddleOCR 官方在线 API（PP-OCRv6），凭据为 AI Studio token，
# 通过环境变量 PADDLEOCR_ACCESS_TOKEN 提供（不落文件、不入 git）。
# 本机不安装任何本地识别模型，推理全部走 PaddleOCR 云端服务。
PADDLEOCR_BASE = "https://paddleocr.aistudio-app.com"


def _paddle_token():
    """返回 PaddleOCR 在线 API token。仅从环境变量 PADDLEOCR_ACCESS_TOKEN 读取。"""
    tok = os.environ.get("PADDLEOCR_ACCESS_TOKEN", "").strip()
    return tok or None


def _paddle_ocr_text(img_bytes, timeout=60):
    """调用 PaddleOCR 在线 API（PP-OCRv6）识别图片，返回识别文本（rec_texts 拼接）。
    流程：multipart 上传 → 轮询 job 状态 → 取 resultUrl.jsonUrl → 解析 rec_texts。
    失败返回 None。"""
    import urllib.request
    token = _paddle_token()
    if not token:
        return None
    jobs_url = PADDLEOCR_BASE + "/api/v2/ocr/jobs"
    boundary = "----sheqi%x" % (int(time.time() * 1000))
    parts = []
    for k, v in (("model", "PP-OCRv6"), ("optionalPayload", "{}")):
        parts.append(("--%s\r\n" % boundary).encode("utf-8"))
        parts.append(("Content-Disposition: form-data; name=\"%s\"\r\n\r\n" % k).encode("utf-8"))
        parts.append((v + "\r\n").encode("utf-8"))
    parts.append(("--%s\r\n" % boundary).encode("utf-8"))
    parts.append(("Content-Disposition: form-data; name=\"file\"; filename=\"captcha.png\"\r\n").encode("utf-8"))
    parts.append(("Content-Type: application/octet-stream\r\n\r\n").encode("utf-8"))
    parts.append(img_bytes)
    parts.append(("\r\n--%s--\r\n" % boundary).encode("utf-8"))
    body = b"".join(parts)
    req = urllib.request.Request(jobs_url, data=body, method="POST")
    req.add_header("Authorization", "Bearer " + token)
    req.add_header("Content-Type", "multipart/form-data; boundary=" + boundary)
    with urllib.request.urlopen(req, timeout=timeout) as r:
        submit = json.loads(r.read().decode("utf-8"))
    job_id = (submit.get("data") or {}).get("jobId")
    if not job_id:
        return None
    for _ in range(20):
        time.sleep(1.5)
        req = urllib.request.Request(jobs_url + "/" + job_id)
        req.add_header("Authorization", "Bearer " + token)
        with urllib.request.urlopen(req, timeout=timeout) as r:
            st = json.loads(r.read().decode("utf-8"))
        d = st.get("data") or {}
        state = d.get("state")
        if state == "done":
            json_url = (d.get("resultUrl") or {}).get("jsonUrl")
            if not json_url:
                return None
            with urllib.request.urlopen(json_url, timeout=timeout) as r:
                jl = r.read().decode("utf-8")
            texts = []
            for line in jl.strip().split("\n"):
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                    for item in (obj.get("result", {}).get("ocrResults") or []):
                        texts.extend((item.get("prunedResult") or {}).get("rec_texts") or [])
                except Exception:
                    continue
            return "".join(texts)
        if state == "failed":
            return None
    return None


def parse_math(raw):
    """解析验证码数学算式文本，返回 (a, op, b, ans) 或 None。
    纯数字答案（如 fixture 固定验证码 7）返回 (n, None, None, n)。"""
    if not raw:
        return None
    raw = raw.strip()
    m = re.search(r"(\d{1,2})\s*([+\-xX*×÷/])\s*(\d{1,2})", raw)
    if not m:
        # 兼容纯数字答案（无运算符）：OCR 识别为单个数字时直接作为答案
        mm = re.search(r"(\d{1,2})", raw)
        if mm:
            n = int(mm.group(1))
            return (n, None, None, n)
        return None
    a, op, b = int(m.group(1)), m.group(2), int(m.group(3))
    if b > 20 and len(m.group(3)) == 2:
        b = int(m.group(3)[0])
    if a > 20 and len(m.group(1)) == 2:
        a = int(m.group(1)[0])
    if a > 20 or b > 20:
        return None
    if op == "+":
        ans = a + b
    elif op == "-":
        if a < b:
            return None
        ans = a - b
    elif op in ("x", "X", "*", "×"):
        ans = a * b
    elif op in ("/", "÷"):
        if b == 0 or a % b != 0:
            return None
        ans = a // b
    else:
        return None
    return (a, op, b, ans)


def ocr_captcha(img_bytes):
    """验证码识别：PaddleOCR 在线 API（多路预处理：原图→灰度2x→二值化），命中算式即返回。
    返回 dict{answer,expr,votes,total,raws} 或 None（失败由上层刷新重试/转人工）。"""
    variants = [("orig", img_bytes)]
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes)).convert("L")
        gray2x = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
        buf = io.BytesIO()
        gray2x.save(buf, format="PNG")
        variants.append(("gray2x", buf.getvalue()))
        for th in (140, 180):
            binary = gray2x.point(lambda x: 0 if x < th else 255)
            buf = io.BytesIO()
            binary.save(buf, format="PNG")
            variants.append(("bin%d" % th, buf.getvalue()))
    except Exception:
        pass
    raw_list = []
    for name, data in variants:
        try:
            text = _paddle_ocr_text(data)
        except Exception as e:
            print("  [PaddleOCR失败] %s" % str(e)[:120])
            continue
        if not text:
            continue
        raw_list.append("%s:%s" % (name, text))
        p = parse_math(text)
        if p:
            expr = ("%s=%s" % (p[0], p[3])) if p[1] is None else ("%s%s%s=%s" % (p[0], p[1], p[2], p[3]))
            return {
                "answer": p[3],
                "expr": expr,
                "votes": 1,
                "total": 1,
                "raws": raw_list,
            }
    return None


def ocr_captcha_protect(img_bytes):
    """识别 12377 云盾保护页字符验证码（PaddleOCR 在线 API，直接取文本）。
    返回去除空白的识别文本；失败返回 None。"""
    try:
        text = _paddle_ocr_text(img_bytes)
    except Exception as e:
        print("  [保护页OCR失败] %s" % str(e)[:120])
        return None
    if not text:
        return None
    text = re.sub(r"\s+", "", text)
    return text or None


def get_captcha_bytes(page):
    """获取表单页验证码图片字节。
    优先 el.screenshot() 截取页面当前显示的图（与提交绑定的验证码一致）；
    page.request 重新请求 src 会生成新验证码，仅作最后兜底。"""
    el = page.query_selector("#verifycode_id")
    if el:
        try:
            return el.screenshot()
        except Exception:
            pass
    src = page.evaluate("const el = document.getElementById('verifycode_id'); el ? el.src : ''")
    if src:
        try:
            resp = page.request.get(src, timeout=20000)
            if resp.status == 200:
                body = resp.body()
                if body and len(body) > 100:
                    return body
        except Exception:
            pass
    return None


def pass_captcha_protection(page, max_try=6, manual=False, ocr_enabled=True):
    """处理 12377 云盾"验证码保护"页（#cap-img / #ans / #submit）。
    OCR 自动识别（PaddleOCR 在线 API）优先，失败转人工输入。
    成功进入真实页面返回 True，否则 False；无保护元素直接返回 True。"""
    for i in range(1, max_try + 1):
        try:
            if not page.query_selector("#cap-img") and not page.query_selector("#ans"):
                print("    [保护页] 无需处理")
                return True
            # 获取验证码图片（data:image 内嵌 base64 或网络图）
            src = page.evaluate("const el = document.getElementById('cap-img'); el ? el.src : ''")
            if not src:
                return False
            if src.startswith("data:"):
                import base64
                img_bytes = base64.b64decode(src.split(",", 1)[1])
            else:
                resp = page.request.get(src)
                img_bytes = resp.body()
            fill_val = None
            if ocr_enabled:
                ans = ocr_captcha_protect(img_bytes)
                print("    [保护页] 第%d次识别: %r" % (i, ans))
                if ans:
                    p = parse_math(ans)
                    if p:
                        fill_val = str(p[3])
                        if p[1] is None:
                            print("    [保护页] 数字 %s → 填 %s" % (p[0], fill_val))
                        else:
                            print("    [保护页] 算式 %s%s%s=? → 填 %s" % (p[0], p[1], p[2], fill_val))
                if fill_val is None:
                    # 多路预处理兜底提高算式命中率
                    try:
                        res = ocr_captcha(img_bytes)
                        if res and res.get("answer") is not None:
                            fill_val = str(res["answer"])
                            print("    [保护页] 多路算式识别: %s → 填 %s" % (res.get("expr", "?"), fill_val))
                    except Exception as e:
                        print("    [保护页] 多路算式识别异常: %s" % str(e)[:80])
                if fill_val is None and ans:
                    # 形态2/3：字母或中文词语 → 填原文
                    fill_val = ans
            if fill_val is None and manual:
                try:
                    fill_val = input("    [保护页] 请查看浏览器中的验证码并输入答案（回车换一张）: ").strip()
                except EOFError:
                    fill_val = ""
            if fill_val is None:
                try:
                    page.click("#update", timeout=3000)
                except Exception:
                    pass
                page.wait_for_timeout(1500)
                continue
            page.fill("#ans", fill_val)
            page.click("#submit", timeout=5000)
            page.wait_for_timeout(3500)
            # 检查是否通过（保护元素消失）
            if not page.query_selector("#cap-img") and not page.query_selector("#ans"):
                print("    [保护页] 已通过")
                return True
            # 仍在保护页：换一张重试
            try:
                page.click("#update", timeout=3000)
            except Exception:
                pass
            page.wait_for_timeout(1500)
        except Exception as e:
            print("    [保护页] 异常: %s" % str(e)[:120])
            try:
                page.wait_for_timeout(1500)
            except Exception:
                pass
    return False


def open_form(page, base_url, manual=False, ocr_enabled=True):
    for attempt in range(1, 4):
        try:
            page.goto(base_url, wait_until="domcontentloaded", timeout=60000)
            print("    [表单] 第%d次 已打开: %s" % (attempt, page.url))
            # 等待 JS 挑战/保护页/表单页任一出现（最多 25 秒）
            deadline = time.time() + 25
            while time.time() < deadline:
                try:
                    if (page.query_selector("#cap-img") or page.query_selector("#isee")
                            or page.query_selector("#entName")):
                        break
                except Exception:
                    pass
                page.wait_for_timeout(1000)
            page.wait_for_timeout(2000)
            # 若为云盾验证码保护页，OCR 自动识别（失败转人工）
            if page.query_selector("#cap-img"):
                ok = pass_captcha_protection(page, manual=manual, ocr_enabled=ocr_enabled)
                if not ok:
                    print("    [表单] 验证码保护页未能通过（第%d次）" % attempt)
            page.wait_for_selector("#isee", state="attached", timeout=15000)
            page.evaluate("document.getElementById('isee').checked = true")
            page.wait_for_timeout(500)
            try:
                page.click("#agree", timeout=10000)
                print("    [表单] 已点击同意")
            except Exception:
                try:
                    page.locator("#agree").click(force=True, timeout=10000)
                    print("    [表单] 强制点击同意")
                except Exception as e:
                    print("    [表单] 点击失败: %s" % str(e)[:120])
        except Exception as e:
            print("    [表单] 准备阶段异常: %s" % str(e)[:150])
        deadline = time.time() + 20
        while time.time() < deadline:
            try:
                ent = page.query_selector("#entName") is not None
            except Exception:
                ent = False
            if ent:
                print("    [表单] 表单页就绪")
                return page
            try:
                ctx = page.context
                for pg in ctx.pages:
                    if "sqjb" in pg.url and pg != page:
                        page = pg
            except Exception:
                pass
            time.sleep(1)
        print("    [表单] 第%d次超时 url=%s pages=%s" % (
            attempt, page.url, [t.url for t in page.context.pages]))
    raise RuntimeError("无法进入举报表单页（#entName 未出现），已重试 3 次")


def select_nice_js(page, select_id, values):
    vals = json.dumps(values if isinstance(values, (list, tuple)) else [values], ensure_ascii=False)
    return page.evaluate("""() => {
        const sel = document.getElementById('%s');
        if (!sel) return 'notfound';
        const want = %s;
        const wrapper = sel.closest('.select-item') || sel.parentElement;
        const nice = wrapper && wrapper.querySelector('.nice-select');
        if (sel.multiple) {
            Array.from(sel.options).forEach(o => { o.selected = want.includes(o.value); });
        } else if (want[0]) {
            sel.value = want[0];
        }
        sel.dispatchEvent(new Event('change', { bubbles: true }));
        if (nice) {
            const cur = nice.querySelector('.current');
            const opt = sel.multiple ? null : Array.from(sel.options).find(o => o.value === sel.value);
            if (cur && opt) cur.textContent = opt.textContent;
            nice.querySelectorAll('li').forEach(li => {
                const v = li.getAttribute('data-value');
                const on = sel.multiple ? want.includes(v) : (v === sel.value);
                li.classList.toggle('selected', !!on);
                if (!sel.multiple && on) li.classList.add('focus');
            });
        }
        return 'ok';
    }""" % (select_id, vals))


def fill_form(page, subject, rec):
    page.fill("#entName", str(subject.get("company_name") or ""))
    for sid, val in (
        ("entType", subject.get("company_type")),
        ("entNature", subject.get("company_nature")),
        ("entIndustry", subject.get("industry")),
        ("entContactType", subject.get("contact_type")),
    ):
        if val:
            select_nice_js(page, sid, [val])
    for fid, val in (
        ("reportor_realname", subject.get("contact_name")),
        ("reportor_telephone", subject.get("contact_phone")),
        ("reportor_email", subject.get("contact_email")),
    ):
        if val and page.query_selector("#" + fid):
            page.fill("#" + fid, str(val))
    if page.query_selector("#targetName"):
        # 举报网站名称 ≤50 字：组内平台去重拼接，超了从尾部逐个删直到可行
        platforms = rec.get("platforms") or [rec.get("platform")]
        page.fill("#targetName", fit_names_to_limit(platforms, 50))
    # 举报账号名称 ≤50 字：组内账号去重拼接，超了从尾部逐个删直到可行
    accounts = rec.get("accounts") or [rec.get("account") or rec.get("platform")]
    page.fill("#targetAccount", fit_names_to_limit(accounts, 50))
    # 合并模式：targetUrl 为 textarea，多 URL 用换行拼接
    urls = rec.get("urls") or ([rec.get("url")] if rec.get("url") else [])
    page.fill("#targetUrl", "\n".join(str(u) for u in urls))
    page.fill("#content", str(rec.get("content") or "")[:CONTENT_MAX])
    select_nice_js(page, "entProofType", [str(rec.get("evidence_type") or "P008")])


def upload_files(page, subject, rec):
    def rel(key):
        return resolve_doc(subject, key)

    jobs = []
    lic = rel("license")
    if lic:
        jobs.append(("upload_file1", [lic]))
    # 卡槽2：身份证正面+反面+授权委托书+在职证明
    ids = [p for p in (rel("id_front"), rel("id_back"), rel("auth"), rel("hand_id"), rel("work_proof")) if p]
    if ids:
        jobs.append(("upload_file2", ids))
    # 举报信：letter/letter_dup/letter2/letter3 全部收集
    # 多份举报信 → 卡槽3 放第一份，卡槽4 放其余；只有一份 → 优先卡槽3（卡槽4 留空）
    letter_files = []
    for k in ("letter", "letter_dup", "letter2", "letter3"):
        p = rel(k)
        if p and os.path.exists(p):
            letter_files.append(p)
    if letter_files:
        if len(letter_files) >= 2:
            jobs.append(("upload_file3", letter_files[:1]))
            jobs.append(("upload_file4", letter_files[1:]))
        else:
            jobs.append(("upload_file3", letter_files))
    for input_id, paths in jobs:
        if len(paths) > 1:
            try:
                page.evaluate("const el = document.getElementById('%s'); if (el) el.multiple = true;" % input_id)
            except Exception:
                pass
        page.set_input_files("#" + input_id, paths)
        time.sleep(1.2)
    return jobs


def dump_form_values(page, rec):
    """dry-run 辅助：提取表单页关键字段当前值（含 select 选中项、上传文件名），返回文本。
    用于无需截图识别的快速核对。"""
    try:
        lines = []
        # 文本输入 / textarea
        for sel, label in (
            ("#entName", "主体名称"),
            ("#reportor_realname", "联系人"),
            ("#reportor_telephone", "联系电话"),
            ("#reportor_email", "邮箱"),
            ("#targetName", "被举报网站/账号名称"),
            ("#targetAccount", "被举报账号"),
            ("#targetUrl", "被举报链接"),
            ("#content", "具体举报内容"),
        ):
            try:
                el = page.query_selector(sel)
                if el:
                    val = el.input_value() if el.get_attribute("type") != "file" else ""
                    lines.append("%s: %s" % (label, val[:2000]))
            except Exception:
                pass
        # select 选中值
        for sel, label in (
            ("#entType", "企业类型"),
            ("#entNature", "企业性质"),
            ("#entContactType", "联系人类型"),
            ("#entProofType", "证据类型"),
        ):
            try:
                el = page.query_selector(sel)
                if el:
                    val = el.evaluate("el => el.options[el.selectedIndex] ? el.options[el.selectedIndex].text : ''")
                    lines.append("%s: %s" % (label, val))
            except Exception:
                pass
        # 文件上传
        for fid in ("upload_file1", "upload_file2", "upload_file3", "upload_file4"):
            try:
                el = page.query_selector("#" + fid)
                if el:
                    names = el.evaluate("el => Array.from(el.files || []).map(f => f.name).join(', ')")
                    lines.append("%s: %s" % (fid, names))
            except Exception:
                pass
        return "\n".join(lines)
    except Exception:
        return None


def refresh_captcha(page):
    try:
        page.evaluate("""() => {
            const el = document.getElementById('verifycode_id');
            if (!el) return;
            const base = (el.src || 'https://new.12377.cn/rpapi/portal/captcha').split('?')[0];
            el.src = base + '?' + Date.now();
        }""")
    except Exception:
        pass
    time.sleep(1.2)


def extract_code(text):
    m = re.search(r"查询码[:：]?\s*([A-Za-z0-9]{8,})", text or "")
    return m.group(1) if m else ""


def make_response_hook(responses):
    """收集提交相关的 POST 响应，并在收到时立即缓存 body（避免后续重复读取报错）。"""
    def hook(resp):
        try:
            if resp.request.method == "POST" and any(k in resp.url for k in ("addsq", "submit", "report")):
                try:
                    resp._cached_body = resp.body()
                except Exception:
                    resp._cached_body = None
                responses.append(resp)
        except Exception:
            pass
    return hook


def wait_result(page, responses, timeout=25000):
    deadline = time.time() + timeout
    last_alerts = []
    while time.time() < deadline:
        url = page.url
        if "jbhz" in url or "success" in url.lower():
            try:
                return True, extract_code(page.inner_text("body")), None, False, False
            except Exception:
                return True, "", None, False, False
        for resp in responses:
            body = getattr(resp, "_cached_body", None)
            if body is None:
                continue
            try:
                data = json.loads(body.decode("utf-8", errors="ignore"))
                s = json.dumps(data, ensure_ascii=False)
            except Exception:
                data, s = None, body.decode("utf-8", errors="ignore")
            if "5000" in s or "超出当天" in s or "举报次数" in s:
                return False, None, s[:120], False, True
            if "验证码" in s:
                # 服务端超限初期可能返回含"验证码"字样的风控提示，需识别上限关键词
                if any(k in s for k in ("上限", "超出", "频繁", "次数已达", "达上限", "5000")):
                    print("    [DEBUG] 服务器返回(疑似上限): %s" % s[:300])
                    return False, None, s[:120], False, True
                print("    [DEBUG] 服务器返回(含验证码): %s" % s[:300])
                return False, None, s[:120], True, False
            code = extract_code(s)
            if code:
                return True, code, None, False, False
            if "成功" in s:
                return True, extract_code(s), None, False, False
        try:
            alerts = page.evaluate("window.__sheqiAlerts || []")
        except Exception:
            alerts = []
        if alerts:
            last_alerts = alerts
            joined = "\n".join(alerts)
            import json as _json
            print("    [DEBUG] alerts=%s" % _json.dumps(alerts, ensure_ascii=False)[:300])
            if "验证码" in joined:
                return False, None, joined[:120], True, False
            if "超出当天" in joined or "举报次数" in joined or "5000" in joined:
                return False, None, joined[:120], False, True
            code = extract_code(joined)
            if code or "成功" in joined:
                return True, code, None, False, False
            # 其余 alert（如"请选择行业分类"）为表单校验失败，立即返回错误，不再空等到超时
            if joined.strip():
                return False, None, joined[:120], False, False
        try:
            body = page.inner_text("body")
            code = extract_code(body)
            if code:
                return True, code, None, False, False
            if "举报成功" in body:
                return True, "", None, False, False
        except Exception:
            pass
        time.sleep(0.5)
    if last_alerts:
        return False, None, "提交结果未知：" + last_alerts[-1][:120], False, False
    return False, None, "提交结果未知（超时）", False, False
def submit_and_detect(page, max_retry=MAX_CAPTCHA_RETRY, ocr_enabled=True, manual=False, responses=None):
    page.evaluate("""() => {
        if (!window.__sheqiAlerts) {
            window.__sheqiAlerts = [];
            window.__sheqiConfirms = [];
            // alert 为真正的错误提示；confirm 为提交确认弹窗，返回 true 继续提交且不混入错误提示
            window.alert = function (m) { window.__sheqiAlerts.push(String(m)); };
            window.confirm = function (m) { window.__sheqiConfirms.push(String(m)); return true; };
        }
    }""")
    for attempt in range(1, max_retry + 1):
        page.evaluate("window.__sheqiAlerts = []")
        if responses:
            responses.clear()  # 避免读到上一次尝试的旧响应导致误判
        answer = None
        if ocr_enabled:
            try:
                img = get_captcha_bytes(page)
                if img:
                    res = ocr_captcha(img)
                    if res:
                        answer = str(res["answer"])
                        print("    [OCR] %s (%d/%d 票)" % (res["expr"], res["votes"], res["total"]))
            except Exception as e:
                print("    [OCR] 识别异常: %s" % e)
        if answer is None:
            if manual:
                try:
                    answer = input("    [人工] 请输入验证码答案（回车取消）: ").strip()
                except EOFError:
                    answer = None
                if not answer:
                    return {"ok": False, "code": None, "error": "用户取消", "captcha_error": False, "daily": False}
            else:
                print("    [OCR] 未能识别，刷新重试 %d/%d" % (attempt, max_retry))
                refresh_captcha(page)
                continue
        page.fill("#verify_code", str(answer))
        page.evaluate("""() => {
            if (typeof submitreportinfo === 'function') { submitreportinfo(); return; }
            const b = document.getElementById('btn_submit') || document.querySelector('button[type=submit], input[type=submit]');
            if (b) b.click();
        }""")
        ok, code, err, cap_err, daily = wait_result(page, responses or [], timeout=25000)
        if ok:
            return {"ok": True, "code": code, "error": None, "captcha_error": False, "daily": False}
        if daily:
            return {"ok": False, "code": None, "error": "今日举报次数已达上限", "captcha_error": False, "daily": True}
        if cap_err:
            print("    [提交] 验证码错误，刷新重试 %d/%d" % (attempt, max_retry))
            refresh_captcha(page)
            continue
        # 有明确错误（如表单校验 alert）立即返回，不再无意义重试
        if err:
            return {"ok": False, "code": None, "error": str(err)[:200], "captcha_error": False, "daily": False}
        if attempt < max_retry:
            refresh_captcha(page)
    return {"ok": False, "code": None, "error": "提交失败或结果未知（多次尝试）", "captcha_error": False, "daily": False}


def run_batch(subject_name=None, dry_run=False, headless=True, limit=0, from_idx=0,
              daily=None, interval=None, manual=False, ocr_enabled=True,
              base_url=DEFAULT_REPORT_URL, record_ids=None,
              cool_every=10, cool_seconds=300, content_file=None):
    ensure_dirs()
    # daily 为 None 表示不设本地软限制，实际以服务端“今日举报次数已达上限”为准
    interval = interval if interval is not None else SUBMIT_INTERVAL
    content_file = content_file or os.path.join(data_dir(), "content_template.txt")
    # 具体举报内容：dry-run 前提示用户填写（≤500字）；未填则自动整合
    if not os.path.exists(content_file):
        try:
            os.makedirs(os.path.dirname(content_file), exist_ok=True)
            with open(content_file, "w", encoding="utf-8") as f:
                f.write("")
            print("== 具体举报内容模板已创建: %s ==" % content_file)
        except Exception as e:
            print("== 创建内容模板失败: %s ==" % e)
    if os.path.exists(content_file):
        with open(content_file, encoding="utf-8-sig") as f:
            tpl = f.read().strip()
        if tpl:
            print("== 将使用用户自定义举报内容（%d 字）: %s ==" % (len(tpl), content_file))
        else:
            print("== 具体举报内容未填写（%s），提交前请补充；不填将自动整合 ==" % content_file)
    subject = get_subject(subject_name)
    if not subject:
        raise ValueError("没有可用主体，请先: python sheqi.py subject add --name <名称>")
    sname = subject.get("name") or subject_name
    missing = subject_missing_files(subject)
    if missing:
        print("[警告] 主体 [%s] 缺少材料: %s（可能无法通过前端校验）" % (sname, ", ".join(missing)))
        if dry_run:
            print("[BLOCKED] dry-run 检测到缺少必需材料，请先补全后再试: %s" % ", ".join(missing))
            raise SystemExit(2)
    # 行业分类必选（12377 表单提交时 getReportorInfo 校验，缺失会导致提交直接失败）
    if not subject.get("industry"):
        print("[警告] 主体 [%s] 未设置行业分类（industry，如 I003=制造业），提交时表单会提示“请选择行业分类”导致失败" % sname)
        if dry_run:
            print("[BLOCKED] dry-run 检测到缺少行业分类，请先补充: python sheqi.py subject set --name %s --industry I003" % sname)
            raise SystemExit(2)

    records = load_records()
    if record_ids:
        pending = [r for r in records if r.get("id") in record_ids]
    else:
        pending = pending_records(records, sname)
    if not pending:
        print("没有待提交记录")
        return

    ok_pending = []
    for r in pending:
        errs = []
        if not str(r.get("url") or "").startswith(("http://", "https://")):
            errs.append("链接无效")
        if len(str(r.get("content") or "")) > CONTENT_MAX:
            errs.append("内容超500字")
        if errs:
            print("  [跳过] %s %s" % (r.get("id"), "; ".join(errs)))
            r["status"] = "failed"
            r["error"] = "; ".join(errs)
        else:
            ok_pending.append(r)
    save_records(records)
    if not ok_pending:
        print("没有可提交的有效记录")
        return

    # 始终合并：日期降序（新→旧）、同一天内 URL 短→长，贪心塞满 500 字符/组，不跨日期
    ok_pending, merge_map = merge_records_to_queue(ok_pending, subject,
                                                   url_max=500, content_file=content_file)
    print("合并投递: %d 条记录 -> %d 次投递（日期新→旧，同日 URL 短→长，贪心塞满 500 字符，不跨日期）" % (
        sum(len(v) for v in merge_map.values()), len(ok_pending)))

    results = load_results()
    today_count = count_today_success(results, sname)
    if daily:
        print("主体: %s | 今日已提交: %d/%d（本地软限制） | 待处理: %d" % (sname, today_count, daily, len(ok_pending)))
        if not dry_run and today_count >= daily:
            print("已达到本地设定每日上限 %d 条，已停止。如需继续，请用 --daily N 调高后重跑。" % daily)
            return
    else:
        print("主体: %s | 今日已提交: %d | 待处理: %d（不设本地上限，以服务端为准）" % (sname, today_count, len(ok_pending)))

    start = min(from_idx, len(ok_pending))
    queue = ok_pending[start:]
    if limit:
        queue = queue[:limit]
    if dry_run and not limit:
        queue = queue[:1]
    print("本次处理 %d 次投递 | 模式: %s | 浏览器: %s" % (
        len(queue), "DRY-RUN（不提交）" if dry_run else "正式提交", "无头" if headless else "可见"))

    daily_stopped = False
    p, browser = launch_browser(headless)
    responses = []
    # 自适应降频：连续 captcha_error 达到阈值时自动拉长间隔并冷却，降低触发服务端验证码风控概率
    captcha_streak = 0
    current_interval = interval
    CAPTCHA_STREAK_LIMIT = 3
    INTERVAL_MAX = 600
    # 主动冷却：每投递 cool_every 组（含失败/异常）后主动歇 cool_seconds 秒，抢在服务端风控前预防（0=关闭）
    cool_streak = 0
    try:
        context = browser.new_context(
            viewport={"width": 1280, "height": 900},
            user_agent=STEALTH_UA,
            locale="zh-CN",
            timezone_id="Asia/Shanghai",
        )
        context.add_init_script(STEALTH_JS)  # 绕过 12377 WAF JS 挑战
        page = context.new_page()
        # 挂在 context 上而非 page：open_form 切到新标签页后仍能捕获提交响应
        context.on("response", make_response_hook(responses))
        for i, rec in enumerate(queue, 1):
            print("\n[%d/%d] %s | %s | %s" % (i, len(queue), rec.get("id"), rec.get("platform"), rec.get("url", "")[:80]))
            rec["retry_count"] = int(rec.get("retry_count") or 0) + 1
            rec["status"] = "submitting"
            # 合并组内原始记录同步置 submitting（mrec 是浅拷贝不在 records 内，
            # 否则 finally 回滚 / recover_stuck 找不到 submitting 的原始记录而失效）
            if merge_map.get(rec.get("id")):
                for rid in merge_map[rec.get("id")]:
                    for rr in records:
                        if rr.get("id") == rid:
                            rr["status"] = "submitting"
                            rr["retry_count"] = int(rr.get("retry_count") or 0) + 1
                            break
            save_records(records)
            try:
                # 提交前动态生成该组举报信（内容含组内全部 URL），并更新主体材料引用
                if merge_map.get(rec.get("id")):
                    try:
                        doc_dir = os.path.join(data_dir(), "documents", sname)
                        os.makedirs(doc_dir, exist_ok=True)
                        lp = generate_letter(subject, rec, out_dir=doc_dir)
                        base_name = "letter_%s" % (rec.get("id") or "merge")
                        t1 = os.path.join(doc_dir, base_name + ".docx")
                        t2 = os.path.join(doc_dir, base_name + "_2.docx")
                        if os.path.abspath(lp) != os.path.abspath(t1):
                            shutil.copy2(lp, t1)
                        shutil.copy2(lp, t2)
                        subject["files"]["letter"] = os.path.relpath(t1, data_dir()).replace("\\", "/")
                        subject["files"]["letter2"] = os.path.relpath(t2, data_dir()).replace("\\", "/")
                    except Exception as e:
                        print("  [警告] 动态生成举报信失败（沿用默认）: %s" % e)
                page = open_form(page, base_url, manual=manual, ocr_enabled=ocr_enabled)
                fill_form(page, subject, rec)
                uploads = upload_files(page, subject, rec)
                print("  [上传] %s" % [(k, len(v)) for k, v in uploads])
                if dry_run:
                    shot = os.path.join(data_dir(), "logs", "dry_run_%s.png" % rec.get("id"))
                    try:
                        page.screenshot(path=shot)
                        print("  [截图] %s" % shot)
                    except Exception:
                        pass
                    try:
                        # 表单值 dump：直接提取关键字段值保存为文本，便于快速核对（无需截图识别）
                        dump = dump_form_values(page, rec)
                        if dump:
                            txt = os.path.join(data_dir(), "logs", "dry_run_%s.txt" % rec.get("id"))
                            with open(txt, "w", encoding="utf-8") as f:
                                f.write(dump)
                            print("  [表单值] %s" % txt)
                    except Exception as e:
                        print("  [表单值] dump 失败: %s" % e)
                    rec["status"] = "pending"
                    rec["error"] = ""
                    save_records(records)
                    print("  [DRY-RUN] 填表+上传完成，未提交")
                    continue
                res = submit_and_detect(page, ocr_enabled=ocr_enabled, manual=manual, responses=responses)
                if res["ok"]:
                    captcha_streak = 0  # 成功一次即清零风控计数
                    rec["status"] = "submitted"
                    rec["query_code"] = res["code"]
                    rec["submit_time"] = now_str()
                    rec["error"] = ""
                    results.append({"id": rec.get("id"), "subject": sname, "success": True,
                                    "code": res["code"], "date": today(), "platform": rec.get("platform"),
                                    "url": rec.get("url"), "time": now_str()})
                    print("  [成功] 查询码: %s" % res["code"])
                else:
                    rec["status"] = "failed"
                    rec["error"] = res["error"]
                    results.append({"id": rec.get("id"), "subject": sname, "success": False,
                                    "error": res["error"], "date": today(), "platform": rec.get("platform"),
                                    "url": rec.get("url"), "time": now_str()})
                    print("  [失败] %s" % res["error"])
                    if res.get("daily"):
                        daily_stopped = True
                        save_records(records)
                        save_results(results)
                        print("今日上限，停止剩余队列")
                        break
                    if res.get("captcha_error"):
                        # 验证码风控：连续达到阈值 → 自动降频 + 冷却
                        captcha_streak += 1
                        if captcha_streak >= CAPTCHA_STREAK_LIMIT:
                            current_interval = min(current_interval * 2, INTERVAL_MAX)
                            cool = current_interval * 2
                            print("  [风控] 连续 %d 次验证码错误，自动降频至 %d 秒/组，冷却 %d 秒..." % (
                                captcha_streak, current_interval, cool))
                            captcha_streak = 0
                            save_records(records)
                            save_results(results)
                            time.sleep(cool)
                # 同步组内原始记录状态（共享查询码/错误）
                if merge_map.get(rec.get("id")):
                    for rid in merge_map[rec.get("id")]:
                        for rr in records:
                            if rr.get("id") == rid:
                                rr["status"] = rec["status"]
                                rr["query_code"] = rec.get("query_code", "")
                                rr["error"] = rec.get("error", "")
                                rr["submit_time"] = rec.get("submit_time", "")
                                break
            except Exception as e:
                rec["status"] = "failed"
                rec["error"] = "%s: %s" % (type(e).__name__, e)
                print("  [异常] %s" % rec["error"])
                if merge_map.get(rec.get("id")):
                    for rid in merge_map[rec.get("id")]:
                        for rr in records:
                            if rr.get("id") == rid:
                                rr["status"] = "failed"
                                rr["error"] = rec["error"]
                                break
                try:
                    shot = os.path.join(data_dir(), "logs", "fail_%s.png" % rec.get("id"))
                    page.screenshot(path=shot)
                    print("  [截图] %s" % shot)
                except Exception:
                    pass
            save_records(records)
            save_results(results)
            # 主动冷却：每投递 cool_every 组（含失败/异常）主动歇 cool_seconds 秒，抢在服务端风控前预防（0=关闭）
            if cool_every > 0 and not dry_run:
                cool_streak += 1
                if cool_streak >= cool_every:
                    print("  [主动冷却] 已投递 %d 组，主动歇 %d 秒预防风控..." % (cool_streak, cool_seconds))
                    time.sleep(cool_seconds)
                    cool_streak = 0
            if i < len(queue) and not dry_run:
                print("  等待 %d 秒..." % current_interval)
                time.sleep(current_interval)
    finally:
        # E-1 兜底：中断/崩溃时把仍卡在 submitting 的记录回滚为 pending，避免永久卡死无法重跑
        for rec in queue:
            if rec.get("status") == "submitting":
                rec["status"] = "pending"
                rec["error"] = "interrupted"
                if merge_map.get(rec.get("id")):
                    for rid in merge_map[rec.get("id")]:
                        for rr in records:
                            if rr.get("id") == rid and rr.get("status") == "submitting":
                                rr["status"] = "pending"
                                rr["error"] = "interrupted"
                                break
        save_records(records)
        try:
            browser.close()
        except Exception:
            pass
        p.stop()
    export_csv(records)
    note = "今日举报次数已达上限" if daily_stopped else None
    xlsx_path = export_results_excel(results, note=note)
    print("\n完成。结果已保存: %s" % RESULTS_FILE)
    if xlsx_path:
        print("Excel 结果（含查询码）: %s" % xlsx_path)
    if daily_stopped:
        print("\n⚠️  服务端提示：今日举报次数已达上限，已自动停止剩余队列。")
        print("   您可以：")
        print("   1) 剩余记录保持 pending 状态，明日直接重跑: python sheqi.py run --headless")
        print("   2) 检查结果 Excel（含查询码）核对今日提交情况")
        print("   3) 如需继续，请等待明日额度刷新后再操作")


def serve_fixtures(port=8765, dirpath=None):
    """启动内置离线模拟表单服务（便携版/冻结版无需 Python 即可自测，开发版需 Python 环境）。"""
    import functools
    import http.server
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    d = dirpath or os.path.join(base, "fixtures")
    if not os.path.isdir(d):
        raise RuntimeError("未找到 fixtures 目录: %s" % d)
    handler = functools.partial(http.server.SimpleHTTPRequestHandler, directory=d)
    srv = http.server.ThreadingHTTPServer(("127.0.0.1", port), handler)
    print("离线模拟表单已启动: http://127.0.0.1:%d/form_test.html （Ctrl+C 停止）" % port)
    try:
        srv.serve_forever()
    except KeyboardInterrupt:
        pass
# ================= 举报信函 =================
def generate_letter(subject, rec, out_dir=None):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    ensure_dirs()
    out_dir = out_dir or os.path.join(data_dir(), "output")
    os.makedirs(out_dir, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("举 报 信")
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.font.bold = True
    doc.add_paragraph("中央网信办（国家互联网信息办公室）违法和不良信息举报中心：")
    doc.add_paragraph("举报人：%s" % (subject.get("company_name") or ""))
    doc.add_paragraph("联系电话：%s" % (subject.get("contact_phone") or ""))
    doc.add_paragraph("邮箱：%s" % (subject.get("contact_email") or ""))
    doc.add_paragraph("")
    # 支持合并模式：组内多条记录（urls/platforms/accounts 列表）
    urls = rec.get("urls") or ([rec.get("url")] if rec.get("url") else [])
    platforms = rec.get("platforms") or ([rec.get("platform")] if rec.get("platform") else [])
    accounts = rec.get("accounts") or ([rec.get("account") or rec.get("platform")] if (rec.get("account") or rec.get("platform")) else [])
    url_txt = "\n".join(str(u) for u in urls) if urls else str(rec.get("url") or "")
    doc.add_paragraph(
        "现举报如下网络信息侵犯我司合法权益，请依法处置：\n\n"
        "平台：%s\n账号：%s\n网址：\n%s\n发布时间：%s\n\n举报内容：\n%s\n"
        % ("、".join(str(x) for x in platforms) or "-",
           "、".join(str(x) for x in accounts) or "-",
           url_txt, rec.get("date") or "", rec.get("content") or "")
    )
    doc.add_paragraph("恳请贵中心核实并依法处理。")
    doc.add_paragraph("")
    doc.add_paragraph("举报人（盖章）：")
    doc.add_paragraph(subject.get("company_name") or "")
    doc.add_paragraph("日期：____年____月____日")
    path = os.path.join(out_dir, "举报信函_%s.docx" % (rec.get("id") or "sample"))
    doc.save(path)
    return path


# ================= 环境准备（国内镜像版，跨电脑部署用） =================
PIP_MIRRORS = [
    "https://mirror.baidu.com/pypi/simple",
    "https://pypi.tuna.tsinghua.edu.cn/simple",
    "https://mirrors.aliyun.com/pypi/simple/",
    "https://pypi.mirrors.ustc.edu.cn/simple/",
]


def _venv_python(venv):
    if os.name == "nt":
        return os.path.join(venv, "Scripts", "python.exe")
    return os.path.join(venv, "bin", "python")


def run_setup(python=None):
    if getattr(sys, "frozen", False):
        print("便携版已内置 Python 环境，无需 setup。")
        return sys.executable
    ensure_dirs()
    venv = os.path.join(data_dir(), ".venv")
    py = python or sys.executable
    venv_py = _venv_python(venv)
    if not os.path.exists(venv_py):
        print("创建虚拟环境: %s" % venv)
        subprocess.check_call([py, "-m", "venv", venv])
    req = os.path.join(os.path.dirname(os.path.abspath(__file__)), "requirements.txt")
    print("升级 pip...")
    subprocess.check_call([venv_py, "-m", "pip", "install", "--disable-pip-version-check",
                           "--upgrade", "pip"])
    last_err = None
    for mirror in PIP_MIRRORS:
        try:
            print("用国内镜像安装依赖: %s" % mirror)
            subprocess.check_call([venv_py, "-m", "pip", "install", "--disable-pip-version-check",
                                   "-i", mirror, "-r", req])
            print("环境就绪。后续命令建议用: %s" % venv_py)
            return venv_py
        except subprocess.CalledProcessError as e:
            last_err = e
            print("镜像 %s 失败，换下一个..." % mirror)
    try:
        print("回退官方源安装...")
        subprocess.check_call([venv_py, "-m", "pip", "install", "--disable-pip-version-check",
                               "-r", req])
        print("环境就绪（官方源）。后续命令建议用: %s" % venv_py)
        return venv_py
    except subprocess.CalledProcessError:
        raise RuntimeError("依赖安装失败：请检查网络/镜像，或离线提供 wheel 包（%s）" % last_err)
